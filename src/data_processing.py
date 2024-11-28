import os
from docx import Document
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from langchain_community.vectorstores import FAISS

# Load data
def read_docx(file_path):
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        print(f"Error reading .docx file {file_path}: {e}")

def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = "\n".join([page.extract_text() for page in reader.pages])
    return text

def load_documents_from_folder(folder_path):
    documents = []
    for root, dirs, files in os.walk(folder_path):  # Sử dụng os.walk để duyệt qua thư mục và các thư mục con
        for filename in files:
            file_path = os.path.join(root, filename)
            try:
                if filename.lower().endswith((".docx", ".pdf", ".txt")):  # Chấp nhận .docx, .pdf, .txt
                    if filename.endswith(".docx"):
                        doc_text = read_docx(file_path)
                        doc_text = doc_text.strip()
                        if doc_text:  # Nếu có nội dung
                            documents.append(doc_text)
                    elif filename.endswith(".pdf"):
                        pdf_text = read_pdf(file_path)
                        pdf_text = pdf_text.strip()
                        if pdf_text:  # Nếu có nội dung
                            documents.append(pdf_text)
                    else:
                        print(f"Unsupported file type: {file_path}")
                        continue
            except Exception as e:
                print(f"Error reading file {filename}: {e}")
    return documents
# print(load_documents_from_folder("/home/minhlahanhne/DATN_pre-trained/data"))
def split_data(texts, chunk_size=100, overlap=50, vector_db_path="./vector_db"):
    model_name = "dangvantuan/vietnamese-embedding"
    model = SentenceTransformer(model_name, device="cpu")  # Sử dụng CPU để debug
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=overlap)
    all_chunks = []
    
    # Tạo đối tượng Document từ mỗi chuỗi trong texts
    documents = [Document(text) for text in texts]
    
    # Chia các văn bản trong documents thành các chunk
    for doc in documents:
        chunks = text_splitter.split_documents([doc])  # Sử dụng split_documents để chia nhỏ văn bản
        all_chunks.extend(chunks)        
    print(f"Total number of chunks: {len(all_chunks)}")
    embeddings = model.encode([chunk.text for chunk in all_chunks], convert_to_tensor=True)
    
    # Save embeddings vector to FAISS
    db = FAISS.from_documents(all_chunks, embeddings)
    db.save_local(vector_db_path)
    print(f"Embeddings saved to {vector_db_path}")
    return db

def main():
    folder_path = "./data"
    documents = load_documents_from_folder(folder_path)
    print(f"Total documents loaded: {len(documents)}")
    
    # Kiểm tra xem thư mục có dữ liệu không
    if not documents:
        print("No documents found in the folder!")
        return
    
    # Chia nhỏ dữ liệu thành chunks và lưu embedding vectors vào FAISS
    split_data(documents)

if __name__ == "__main__":
    main()
